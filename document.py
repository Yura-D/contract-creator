from docx import Document
from docx.shared import Inches
import gspread
from oauth2client.service_account import ServiceAccountCredentials
from datetime import datetime
import os
import configuration
import gdrive_api
from re import sub

scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
creds = ServiceAccountCredentials.from_json_keyfile_name("client_secret.json", scope)
client = gspread.authorize(creds)


sheet = client.open(configuration.sheet_name).sheet1

title_name = sheet.findall("ПІБ як в паспорті")
if len(title_name) > 1:
    print("There are more than one cell with - ПІБ як в паспорті:\n", title_name)
    print("Please change it in Google sheet")
    quit()
else: 
    title_name = str(title_name)
# checker that title on it place

parts_title_name = title_name.split()
part_pos = parts_title_name[1]
c_position = part_pos.find("C")
number_column_name = part_pos[c_position+1:]


### Work with templates ###

# You can put above the names of template files and it will work with it

print("Templates: ")
templates = gdrive_api.get_list(configuration.templates_folder)
for template in templates:
    print(template, "-", templates[template][0])

# for printing the list of all templates and create the dict with this tampletes and number of choosing


print("\nPlease choose the files you want to fill. Use \",\" if you want few documents (Example: \"1, 5, 12\").")
get_file = input("Choose the files: ")
choose = sub('[\s]', '', get_file)

templates_choose = list() # make template choose for google folder list and download the choose list

for number_choose in choose.split(','):
   gdrive_api.doc_download(templates[int(number_choose)][1], 
                            templates[int(number_choose)][0]+
                            '.docx', 
                            'temp'+ os.sep)
   # templates_choose.append(templates_dict.get(number_choose.strip()))
# for choosing templates



templates_choose = os.listdir("temp/")
templates_choose.remove(".gitignore")

### To get contract date ###

while True:
    contract_date = input("\nPlease type the date of contract - dd.mm.yyyy (Example - 08.05.2018):\n")
    try:
        contract_datetime = datetime.strptime(contract_date, "%d.%m.%Y")
        if contract_datetime.year < 2000 or contract_datetime.year > 2050:
            print("Incorrect format")
            continue
    except ValueError:
        print("Incorrect format")
        continue
    break
# checker to corect date format

def get_month(contract_date):
    return {
        "01": "січня",
        "02": "лютого",
        "03": "березня",
        "04": "квітня",
        "05": "травня",
        "06": "червня",
        "07": "липня",
        "08": "серпня",
        "09": "вересня",
        "10": "жовтня",
        "11": "листопада",
        "12": "грудня"
    }.get(contract_date[3:5])


month_by_word = get_month(contract_date)
date_by_word = contract_date[0:2] + " " + month_by_word + " " + contract_date[6:10]
# date by word for replacing date instide in contract


### Search ###

column_names = sheet.col_values(number_column_name)
search_results = list()

search = input("\nChoose the employees that you want to have. If want several employees use \",\".\n\nFind the emploees: ")
search_names = search.split(",")
clipboard = list()

for s_name in search_names:
    while True:
        for c_name in column_names[1:]:
            if s_name.strip() in c_name:
                clipboard.append(c_name)
            else: pass
        
        if len(clipboard) > 1:
            print("\nThere is more than one result: ", s_name)
            for cb in clipboard:
                print("-", cb)
            print("\nIf you want to pass, please type: '0'")
            s_name = input("\nPlease Try again: ")
            del clipboard[:]
        elif len(clipboard) == 0:
            print("\nWe don't found: ", s_name)
            print("\nIf you want to pass, please type: '0'")
            s_name = input("\nPlease Try again: ")
            del clipboard[:]
        else:
            search_results.append(clipboard[0])
            del clipboard[:]
            break

        if s_name == "0":
            break
        else: pass


def text_replace(old_text, new_text, file):
    doc = Document(file)
    
    
    def replace_part(old_text, new_text, p):
        if old_text in p.text:
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for line in inline:
                if old_text in line.text:
                    text = line.text.replace(old_text, new_text)
                    line.text = text

    for p in doc.paragraphs:
        replace_part(old_text, new_text, p)
       
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_part(old_text, new_text, p)
                
    doc.save(new_file)


### To create contracts ###
print("\nList of emploees:")
for r_name in search_results:
    print("- " + r_name)

    row_number = 0
    for name in column_names:
        row_number = row_number + 1
        if name == r_name:
            break

    unit = sheet.row_values(row_number)

    if len(unit[8]) < 1:
        fop_address = unit[7]
    else:
        fop_address = unit[8]
    # for using Address if there not FOP Address

    name_split = r_name.split()
    name_initials = name_split[0] + " " + name_split[1][0] + "." + name_split[2][0] + "."
    #for making name with initials







    dir_path = "ready_to_print/" + r_name + "/"
    if not os.path.exists(dir_path):
        os.makedirs(dir_path)
    # for creating the directorys

    for template in templates_choose:
        template_path = "temp/" + template
        file_name_date = contract_date[6:10] + contract_date[2:6] + contract_date[0:2]
        new_file = dir_path + file_name_date + " - " + str(template)
        # to create the file name with revers position of date

        text_replace("${Contract date}",        date_by_word, template_path)
        text_replace("${Name}",                 r_name, new_file)
        text_replace("${Passport ID}",          unit[3], new_file)
        text_replace("${Passport ID letter}",   unit[3][0:2], new_file)
        text_replace("${Passport ID number}",   unit[3][2:], new_file)
        text_replace("${Passport address}",     unit[4], new_file)
        text_replace("${Passport date}",        unit[6], new_file)
        text_replace("${Address}",              unit[7], new_file)
        text_replace("${FOP address}",          fop_address, new_file)
        text_replace("${ID}",                   unit[10], new_file)
        text_replace("${Bank}",                 unit[11], new_file)
        text_replace("${Bank info}",            unit[13], new_file)
        text_replace("${Person bank info}",     unit[14], new_file)
        text_replace("${FOP ID}",               unit[15], new_file)
        text_replace("${FOP ID date}",          unit[16], new_file)
        text_replace("${Born}",                 unit[20], new_file)
        text_replace("${Name initials}",        name_initials, new_file)
    # to replace everything that you need in the template

print("\nDone")